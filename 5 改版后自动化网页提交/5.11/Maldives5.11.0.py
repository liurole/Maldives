# -*- coding: utf-8 -*-
"""
Created on Wed Feb 14 22:03:08 2018

@author: Se

下载chromedrive驱动 
使用Selenium需要选择一个调用的浏览器并下载好对应的驱动，本文使用chrome浏览器，当然也可以用FireFox等
http://www.seleniumhq.org/download/ 找到Google Chrome Driver链接
对应驱动放在python目录下面的scripts目录中，例如C:\ProgramData\Anaconda3\envs\python35\Scripts
"""
import requests
import re
import xlrd
import urllib
import hashlib
import random
# 可直接通过pip install 安装
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docxtpl import DocxTemplate
from bs4 import BeautifulSoup

def open_excel(file):  
    data = xlrd.open_workbook(file)  
    return data

def excel_table_byindex(file, colnameindex = 0, by_index = 0):  
    data = open_excel(file)  
    table = data.sheets()[by_index]  
    nrows = table.nrows #行数   
    colnames = table.row_values(colnameindex) #某一行数据  
    list = []  
    for rownum in range(1,nrows):  
        row = table.row_values(rownum)#以列表格式输出  
        if row:  
            app = {}  
            for i in range(len(colnames)):  
                app[colnames[i]] = row[i]  
            list.append(app)#向列表中插入字典类型的数据  
    return list 

def baidu_fanyi(q: str, fro: str = 'auto', to: str = 'zh', timeout: int = 5):
    """
    调用百度翻译API实现在线翻译
    :param q:请求翻译query
    :param fro:翻译源语言，默认自动识别
    :param to:译文语言，默认zh
    :param timeout:设置超时时间，默认5秒
    :return:正常返回json格式数据，否则返回None，结果保存在['trans_result'][0]['dst']中
    """
    appid = '20171203000101916'  # fill in your app_id
    salt = str(random.random())
    key = '8W9VbfBIzhNOhYT729WW'  # fill in your key
    text = appid + q + salt + key
    md5 = hashlib.md5()
    md5.update(text.encode('UTF-8'))
    sign = md5.hexdigest()
    url = 'https://fanyi-api.baidu.com/api/trans/vip/translate'
    argv = {
        'appid': appid,
        'salt': salt,
        'key': key,
        'sign': sign,
        'q': q,
        'from': fro,
        'to': to
    }
    try:
        r = requests.get(url, params=argv, timeout=timeout)
    except:
        return None
    return r.json()

def youdao_fanyi(q: str, fro: str = 'auto', to: str = 'zh', timeout: int = 5):
    """
    调用有道翻译API实现在线翻译
    :param q:请求翻译query
    :param fro:翻译源语言，默认自动识别
    :param to:译文语言，默认zh
    :param timeout:设置超时时间，默认5秒
    :return:正常返回json格式数据，否则返回None，结果保存在['translation']中
    """
    appid = '68a68ddaf01f7b2f'  # fill in your app_id
    salt = str(random.random())
    key = 'AQwzCoZjhpMNVZApqUb9PWEghKv9eBLD'  # fill in your key
    text = appid + q + salt + key
    md5 = hashlib.md5()
    md5.update(text.encode('UTF-8'))
    sign = md5.hexdigest()
    url = 'https://openapi.youdao.com/api/'

    argv = {
        'appKey': appid,
        'q': q,
        'from': fro,
        'to': to,
        'salt': salt,
        'sign': sign
    }
    try:
        r = requests.get(url, params=argv, timeout=timeout)
    except:
        return None
    return r.json()

# 利用request，导入cookies，header进行关键词网页搜索，选择第二栏
def check(url):
    headers = {
    'Accept': 'text/html, application/xhtml+xml, image/jxr, */*',
    'Accept-Encoding': 'gzip, deflate',
    'Accept-Language': 'zh-CN',
    'Connection': 'keep-alive',
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/51.0.2704.79 Safari/537.36 Edge/14.14393'
    }
    try:
        response = requests.get(url, headers = headers)
        if response.status_code == 200:
            return response.text
        return None
    except ConnectionError:
        print('Error occurred')
        return None

# 保存所有页面
def search_all_urls(url):
    response = check(url)
    soup = BeautifulSoup(response, 'lxml')
    title = soup.select('#hs_cos_wrapper_name')
    temp = title[0].get_text()
    temp_baidu = baidu_fanyi(temp)
    temp_baidu = temp_baidu['trans_result'][0]['dst']
    temp_youdao = youdao_fanyi(temp)
    temp_youdao = temp_youdao['translation']
    doc = DocxTemplate('templet.docx')
    context = {
        'MyTitle1' : temp ,
        'MyTitle2' : temp_baidu ,
        'MyTitle3' : temp_youdao
    }
    
    file_name = re.sub('[^A-Za-z\s]','',temp) + '.docx'
    doc.render(context)
    doc.save("generated_temp.docx")
    
    paras = soup.select('#hs_cos_wrapper_post_body p')
    for i in paras:
        temp = i.get_text()
        temp_baidu = baidu_fanyi(temp)
        temp_baidu = temp_baidu['trans_result'][0]['dst']
        temp_youdao = youdao_fanyi(temp)
        temp_youdao = temp_youdao['translation']
        
        doc.add_paragraph(temp)
        doc.add_paragraph('--------------------------------------------------')
        doc.add_paragraph(temp_baidu)
        doc.add_paragraph('--------------------------------------------------')
        doc.add_paragraph(temp_youdao)   
        doc.add_paragraph('')

    doc.save(file_name)              
    print(file_name + '------Done!')

if __name__ == '__main__':
    
    tables = excel_table_byindex('刘晨-主题提交 - 2月11日.xlsx', 0, 1) 
    
    for i in tables:
        url = i['参考链接']
        search_all_urls(url)

    
    