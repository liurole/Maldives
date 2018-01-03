# -*- coding: utf-8 -*-
"""
Created on Fri Dec 29 08:03:29 2017

@author: liurole

马尔代夫V5.1.1

实现了文章的自动化网页生成同时word本地备份，防止撕逼

需指定生成index（-i xxx）

配置文件：
templet.docx 模板文件
paragraph.csv 库函数文件，此处更新库语句
detail.csv 特性文件，此文件为所有产品的信息集合

定义一下各文件格式：
templet.docx 变数不大，这个到时候我来改就行
paragraph.csv 这个是变数最大的文件，考虑到语句数量可能不同，这里每一句用的是行排列
detail.csv 这里加一个表头（第一行），方便调用，同时添加一列序号，方便每次生成

"""

# 这些都是配置库，不用管
import sys
import getopt
import csv
import codecs
import random
# 可直接通过pip install 安装
from docx import Document
from docxtpl import DocxTemplate


# 帮助run Maldives5.2.5.py --help
if __name__ == '__main__':
    opts, args = getopt.getopt(sys.argv[1:], 'hi:', [ 'help', 'index=' ])
    
    index = 1
    err = []
    
    # 入口函数，不明白怎么调用参数的可以看下
    for key, value in opts:
        if key in ['-h', '--help']:
            print('马尔代夫V5.2.5')
            print('参数定义：')
            print('-h, --help\t显示帮助')
            print('-i, --index\t序号')
            sys.exit(0)
        if key in ['-i', '--index']:
            index = value
    print('马尔代夫V5.2.5 function with num: ', index ) 

# STEP 1，首先读取所需的paragraph.csv，detail.csv文件，被注释的两端分别用来检测是否运行成功

    # 查找所选index，这里利用变量pick保存（字典类型）
    with open("detail.csv", 'r') as f:
        reader = csv.DictReader(f)
        for row in reader:
            if row['序号'] == str(index):
                pick = row
        f.close()
                
#    # 测试输出，查看是否正确，讲道理不会出问题的，看雇的测试员怎么说
#    print(pick)

    # 读取本地库，row是list表，我们统计有数据的列，利用temp_num存储每行数量
    # 利用temp_val存储每行的具体值，利用ref_num,ref_val存储整个paragraph.csv文件
    ref_num = []
    ref_val = []
    with open("paragraph.csv", 'r') as f:
        reader = csv.reader(f)
        for row in reader:
            temp_num = 0
            temp_val = []
            for i, val in enumerate(row):
                if i == 0:
                    pass
                elif val != '':
                    temp_num += 1
                    temp_val.append(val)
            ref_num.append(temp_num)
            ref_val.append(temp_val)    
        f.close()
        
    ref_app = []
    with codecs.open('app.csv', "r", 'utf_8_sig') as f:
        reader = csv.reader(f)
        for row in reader:
            ref_app.append(row[0])    
        f.close()
        
#    # 测试是否储存成功        
#    for i, temp_num in enumerate(ref_num):
#        for j in ref_val[i]:
#            print(j)    

# STEP 2，三段话的随机组合生成

#    # 伪随机数生成模块。如果不提供 seed，默认使用系统时间。使用相同的 seed，可以获得完全相同的随机数序列，常用于算法改进测试。
#    random.seed(10086)

    # list 格式的三段话
    pl1 = []
    pl2 = []
    pl3 = []        

    # 生成三段话
    for i, temp_num in enumerate(ref_num):
        if i < 4:
            if i == 1:
                if pick['分类'] == 'TFT Digital':
                    pl1.append(ref_val[i][0])
                elif pick['分类'] == 'Automotive TFT':
                    pl1.append(ref_val[i][1])
                elif pick['分类'] == 'Monochrome-Character':
                    pl1.append(ref_val[i][2])
                elif pick['分类'] == 'Monochrome-Graphic':
                    pl1.append(ref_val[i][3])
            else:
                r1 = random.randint(0, ref_num[i] - 1)
                pl1.append(ref_val[i][r1])
        elif i < 9:
            if i == 5:
                if pick['类型'] == 'Transflective':
                    pl2.append(ref_val[i][0])
                elif pick['类型'] == 'Transmissive':
                    pl2.append(ref_val[i][1])
            else:
                r2 = random.randint(0, ref_num[i] - 1)
                pl2.append(ref_val[i][r2])
        else:
            if i == 10:
                if pick['介质'] == 'LVDS':
                    pl3.append(ref_val[i][0])
                elif pick['介质'] == 'CMOS':
                    pl3.append(ref_val[i][1])
                elif 'RGB' in pick['介质']:
                    pl3.append(ref_val[i][2])
            else:                    
                r3 = random.randint(0, ref_num[i] - 1)
                pl3.append(ref_val[i][r3])
            
#    # 测试输出结果        
#    print(pl1)
#    print(pl2)
#    print(pl3)

# STEP 3，参数的替换
    
    first = random.randint(0, len(ref_app) - 1)
    second = random.randint(0, len(ref_app) - 1)
    while first == second:
        second = random.randint(0, len(ref_app) - 1)

    temp1 = pl1[0].replace('XXX', pick['型号'])
    
    temp1 = temp1.replace('YYY', pl1[1])
    p1 = temp1 + pl1[2] + pl1[3]

    temp1 = pl2[0].replace('XXX', pick['型号'])
    if pick['大小'] == '':
        err.append('大小')
    else:
        temp1 = temp1.replace('YYY', pick['大小'])
    if pick['分辨率'] == '':
        err.append('分辨率')
    else:
        temp1 = temp1.replace('ZZZ', pick['分辨率'])
    if pick['背光'] == '':
        err.append('背光')
    else:
        temp3 = pl2[2].replace('XXX', pick['背光'])
    p2 = temp1 + pl2[1] + temp3 + pl2[3] + pl2[4]

    if pick['介质'] == '':
        err.append('介质')
    else:
        temp1 = pl3[0].replace('XXX', pick['介质'])
    if pick['亮度'] == '':
        err.append('亮度')
    else:
        temp1 = temp1.replace('YYY', pick['亮度'])
    if pick['对比度'] == '':
        err.append('对比度')
    else:
        contrast = pick['对比度']
        contrast_list = contrast.split(':', 1)
        contrast = contrast_list[0] + ':1'
        temp1 = temp1.replace('ZZZ', contrast)
    
    if pick['Operating Temperature (oC)'] == '':
        err.append('操作温度')
    else: 
        temp2 = pl3[2].replace('XXX', pick['Operating Temperature (oC)'])
    if pick['Storage Temperature (oC)'] == '':
        err.append('储存温度')
    else:
        temp2 = temp2.replace('YYY', pick['Storage Temperature (oC)'])
    if pick['Outer Dimension (mm):'] == '':
        err.append('尺寸')
    else:
        temp_d = pick['Outer Dimension (mm):']
        temp_d = temp_d.replace(' ', '')
        temp_d = temp_d.replace('x', ' X ')
        temp3 = pl3[3].replace('XXX', temp_d)
    if pick['Mass (g):'] == '':
        err.append('重量')
    else:
        temp3 = temp3.replace('YYY', pick['Mass (g):'])
    p3 = temp1 + pl3[1] + temp2 + temp3 + pl3[4]
    err.append('视角')
    
#    # 测试输出结果 
#    print(p1)
#    print(p2)
#    print(p3)

# STEP 4，word的保存

    # 标题由detail导入
    MyContent_type = '新产品'
    MyContent_abstract = ''  # 摘要默认为空，手写
    MyContent_factory = 'Kyocera(京瓷)'
    MyContent_device = '显示屏，液晶显示屏，薄膜晶体管，TFT， Thin Film Transistor'
    # 型号由detail导入
    MyContent_app = ''  # 应用选用随机生成
    
    MyContent_key = '薄膜晶体管，TFT，液晶显示屏，显示屏，' + pick['分辨率']
    if pick['大小'] != "":
        MyContent_key += '，' + pick['大小'] + '英寸'
        
    MyContent_name = '刘阳（翻译）'
    MyContent_author = '穿山甲说'
    # 参考链接由detail导入

    doc = DocxTemplate("templet.docx")
    context = {
        'MyTable' : [
            {'title' : '【产品】' + pick['标题'], 'type' : MyContent_type, 'abstract' : MyContent_abstract,
             'factory' : MyContent_factory, 'device' : MyContent_device, 'version' : pick['型号'], 
             'app' : MyContent_app, 'key' : MyContent_key, 'name' : MyContent_name, 
             'author' : MyContent_author, 'ref' : pick['网址'] + '（部分网络需要VPN才能访问）' }
        ],
        'MyFirstP' : p1 ,
        'MySecondP' : p2 ,
        'MyThirdP' : p3
    }
    
    doc.render(context)
    doc.save("generated_temp.docx")
    
    doc = Document("generated_temp.docx")
    
    p1 = doc.add_paragraph('')
    r1 = p1.add_run(pick['型号'] + '的主要特点：')
    
    if pick['特性1'] != '':
        paragraph = doc.add_paragraph(pick['特性1'])
        paragraph_format = paragraph.paragraph_format
    if pick['特性2'] != '':
        paragraph = doc.add_paragraph(pick['特性2'])
        paragraph_format = paragraph.paragraph_format
    
    doc.add_paragraph('')
    p2 = doc.add_paragraph('')
    r2 = p2.add_run(pick['型号'] + '的典型应用：')
    
    application = pick['应用']
    application = application.replace('?', '')
    app_list = application.split(',')
    for i in app_list:
        i = i.strip()
        paragraph = doc.add_paragraph(i)  
        paragraph_format = paragraph.paragraph_format
    doc.add_paragraph(ref_app[first]) 
    doc.add_paragraph(ref_app[second]) 
    
    file_name = 'result_temp.docx'
    
    
    doc.save(file_name)  
