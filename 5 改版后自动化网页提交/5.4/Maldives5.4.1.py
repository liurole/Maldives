# -*- coding: utf-8 -*-
"""
Created on Fri Dec 29 08:03:29 2017

@author: liurole

马尔代夫V5.4.1

针对于新电元快恢复二极管

实现了文章的自动化网页生成同时word本地备份，防止撕逼

需指定生成index（-i xxx）

配置文件：
templet.docx 模板文件
paragraph.csv 库函数文件，此处更新库语句
detail.csv 特性文件，此文件为所有产品的信息集合
app.csv 应用文件

定义一下各文件格式：
templet.docx 变数不大，这个到时候我来改就行
paragraph.csv 这个是变数最大的文件，考虑到语句数量可能不同，这里每一句用的是行排列
detail.csv 这里加一个表头（第一行），方便调用，同时添加一列序号，方便每次生成

"""

# 这些都是配置库，不用管
import sys
import getopt
import csv
import codecs
import random
import xlrd  

# 可直接通过pip install 安装
from docx import Document
from docxtpl import DocxTemplate

def open_excel(file):  
    data = xlrd.open_workbook(file)  
    return data

def excel_table_byindex(file, colnameindex = 0, by_index = 0):  
    data = open_excel(file)  
    table = data.sheets()[by_index]  
    nrows = table.nrows #行数   
    colnames = table.row_values(colnameindex) #某一行数据  
    list = []  
    for rownum in range(1,nrows):  
        row = table.row_values(rownum)#以列表格式输出  
        if row:  
            app = {}  
            for i in range(len(colnames)):  
                app[colnames[i]] = row[i]  
            list.append(app)#向列表中插入字典类型的数据  
    return list  

# 帮助run Maldives5.1.1.py --help
if __name__ == '__main__':
    opts, args = getopt.getopt(sys.argv[1:], 'hi:', [ 'help', 'index=' ])
    
    index = 1
    
    # 入口函数，不明白怎么调用参数的可以看下
    for key, value in opts:
        if key in ['-h', '--help']:
            print('马尔代夫V5.4')
            print('参数定义：')
            print('-h, --help\t显示帮助')
            print('-i, --index\t序号')
            sys.exit(0)
        if key in ['-i', '--index']:
            index = value
    print('马尔代夫V5.4 function with num: ', index ) 

# STEP 1，首先读取所需的paragraph.csv，detail.csv文件，被注释的两端分别用来检测是否运行成功

    # 查找所选index，这里利用变量pick保存（字典类型）
    tables = excel_table_byindex('detail.xlsx', 0, 3) 
    pick_temp = tables[int(index) - 1]
    pick = {}
    
    for key in pick_temp:
        pick[key] = str(pick_temp[key])
         
    ref_app = []
    with codecs.open('app.csv', "r") as f:
        reader = csv.reader(f)
        for row in reader:
            ref_app.append(row[0])    
        f.close()
                
#    # 测试输出，查看是否正确，讲道理不会出问题的，看雇的测试员怎么说
#    print(pick)

    # 读取本地库，row是list表，我们统计有数据的列，利用temp_num存储每行数量
    # 利用temp_val存储每行的具体值，利用ref_num,ref_val存储整个paragraph.csv文件
    ref_num = []
    ref_val = []
    with codecs.open('paragraph.csv', "r") as f:
        reader = csv.reader(f)
        for row in reader:
            temp_num = 0
            temp_val = []
            for i, val in enumerate(row):
                if i == 0:
                    pass
                elif val != '':
                    temp_num += 1
                    temp_val.append(val)
            ref_num.append(temp_num)
            ref_val.append(temp_val)    
        f.close()
        
#    # 测试是否储存成功        
#    for i, temp_num in enumerate(ref_num):
#        for j in ref_val[i]:
#            print(j)    

# STEP 2，三段话的随机组合生成

#    # 伪随机数生成模块。如果不提供 seed，默认使用系统时间。使用相同的 seed，可以获得完全相同的随机数序列，常用于算法改进测试。
#    random.seed(10086)

    # list 格式的三段话
    pl1 = []
    pl2 = []
    pl3 = []        

    # 生成三段话
    for i, temp_num in enumerate(ref_num):
        if i < 3:
            r1 = random.randint(0, ref_num[i] - 1)
            pl1.append(ref_val[i][r1])
        elif i < 6:
            r2 = random.randint(0, ref_num[i] - 1)
            pl2.append(ref_val[i][r2])   
        else:
            r3 = random.randint(0, ref_num[i] - 1)
            pl3.append(ref_val[i][r3])
            
#    # 测试输出结果        
#    print(pl1)
#    print(pl2)
#    print(pl3)

# STEP 3，参数的替换
    r_list = range(0, len(ref_app))
    app_index = random.sample(r_list, 3)

    temp1 = pl1[0].replace('XXX', pick['型号'])
    temp2 = pl1[1].replace('XXX', pick['VRRM'])
    temp2 = temp2.replace('YYY', pick['IF'])
    p1 = temp1 + temp2 + pl1[2]
    
    temp1 = pl2[0].replace('YYY', pick['型号'])
    temp1 = temp1.replace('XXX', pick['VFMAX'])
    temp2 = pl2[1].replace('XXX', (pick['IR']))
    temp3 = pl2[2].replace('XXX', pick['IFSM'])
    p2 = temp1 + temp2 + temp3
    
    temp1 = pl3[0].replace('ZZZ', pick['型号'])
    temp1 = temp1.replace('XXX', pick['Type'])
    out_size = pick['Size']
    out_size = out_size.replace('x', ' X ')
    temp1 = temp1.replace('YYY', (out_size + '（单位mm）'))
    temp2 = pl3[1].replace('XXX', pick['TRR'])
    temp2= temp2.replace('YYY', pick['结温'])
    temp_s = pick['保存温度']
    temp_s = temp_s.replace('to', '至')
    temp2= temp2.replace('ZZZ', temp_s)
    p3 = temp1 + temp2
    
#    # 测试输出结果 
#    print(p1)
#    print(p2)
#    print(p3)

# STEP 4，word的保存

    # 标题由detail导入
    MyContent_type = '新产品'
    MyContent_abstract = ''  # 摘要默认为空，手写
    MyContent_factory = 'ShinDengen(新电元)'
    MyContent_device = '二极管，快恢复二极管，Super Fast Recovery Diodes'
    # 型号由detail导入
    MyContent_app = ''  # 应用选用随机生成
    MyContent_key = '反向恢复时间，最大反向电压，最大正向整流电流，最大正向电压，正向平均整流电流，最大正向导通电压，正向峰值浪涌电流'
    MyContent_name = '刘阳（翻译）'
    MyContent_author = '穿山甲说'
    # 参考链接由detail导入

    doc = DocxTemplate("templet.docx")
    my_title = '图1：' + pick['型号'] + '封装示意图'
    context = {
        'MyTable' : [
            {'title' : '【产品】' + pick['标题'], 'type' : MyContent_type, 'abstract' : MyContent_abstract,
             'factory' : MyContent_factory, 'device' : MyContent_device, 'version' : pick['型号'], 
             'app' : MyContent_app, 'key' : MyContent_key, 'name' : MyContent_name, 
             'author' : MyContent_author, 'ref' : '' }
        ],
        'MyFirstP' : p1 ,
        'MySecondP' : p2 ,
        'MyThirdP' : p3 ,
        'MyTitle' : my_title
    }
    
    doc.render(context)
    doc.save("generated_temp.docx")
    
    doc = Document("generated_temp.docx")
    
    p1 = doc.add_paragraph('')
    r1 = p1.add_run(pick['型号'] + '的主要特点：')
    
    application = pick['特性']
    app_list = application.split('•')
    index = 0
    for i in app_list:
        if index == 0:
            index = 1
            continue
        else:
            i = i.strip()
            paragraph = doc.add_paragraph('• ' + i)  
            paragraph_format = paragraph.paragraph_format
    
    f1 = '• 反向电压最大为' + pick['VRRM'] + 'V'
    f2 = '• 平均正向整流电流为' + pick['IF'] + 'A'
    f3 = '• 可承受峰值正向浪涌电流达' + pick['IFSM'] + 'A'
    f4 = '• 结温为' + pick['结温'] + '℃，保存温度为' + pick['保存温度'] + '℃'
    f5 = '• 反向恢复时间' + pick['TRR'] + 'ns'
    f6 = '• 小型' + pick['Type'] +'封装'
    f7 = '• 符合AEC-Q101标准'
    doc.add_paragraph(f1)
    doc.add_paragraph(f2)
    doc.add_paragraph(f3)
    doc.add_paragraph(f4)
    doc.add_paragraph(f5)
    doc.add_paragraph(f6)
    doc.add_paragraph(f7)   
    doc.add_paragraph('')
    
    p2 = doc.add_paragraph('')
    r2 = p2.add_run(pick['型号'] + '的典型应用：')
    
    doc.add_paragraph(ref_app[app_index[0]]) 
    doc.add_paragraph(ref_app[app_index[1]]) 
    doc.add_paragraph(ref_app[app_index[2]]) 
    
    file_name = 'result_temp.docx'
    
    doc.save(file_name)  
