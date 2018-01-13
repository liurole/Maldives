# -*- coding: utf-8 -*-
"""
Created on Fri Dec 29 08:03:29 2017

@author: liurole

马尔代夫V5.5.1

针对于新电元快恢复二极管

实现了文章的自动化网页生成同时word本地备份，防止撕逼

需指定生成index（-i xxx）

配置文件：
templet.docx 模板文件
paragraph.csv 库函数文件，此处更新库语句
detail.csv 特性文件，此文件为所有产品的信息集合
app.csv 应用文件

定义一下各文件格式：
templet.docx 变数不大，这个到时候我来改就行
paragraph.csv 这个是变数最大的文件，考虑到语句数量可能不同，这里每一句用的是行排列
detail.csv 这里加一个表头（第一行），方便调用，同时添加一列序号，方便每次生成

"""

# 这些都是配置库，不用管
import sys
import getopt
import csv
import codecs
import random
import xlrd  

# 可直接通过pip install 安装
from docx import Document
from docxtpl import DocxTemplate

def open_excel(file):  
    data = xlrd.open_workbook(file)  
    return data

def excel_table_byindex(file, colnameindex = 0, by_index = 0):  
    data = open_excel(file)  
    table = data.sheets()[by_index]  
    nrows = table.nrows #行数   
    colnames = table.row_values(colnameindex) #某一行数据  
    list = []  
    for rownum in range(1,nrows):  
        row = table.row_values(rownum)#以列表格式输出  
        if row:  
            app = {}  
            for i in range(len(colnames)):  
                app[colnames[i]] = row[i]  
            list.append(app)#向列表中插入字典类型的数据  
    return list  

# 帮助run Maldives5.5.1.py --help
if __name__ == '__main__':
    opts, args = getopt.getopt(sys.argv[1:], 'hi:', [ 'help', 'index=' ])
    
    index = 1
    
    # 入口函数，不明白怎么调用参数的可以看下
    for key, value in opts:
        if key in ['-h', '--help']:
            print('马尔代夫V5.5')
            print('参数定义：')
            print('-h, --help\t显示帮助')
            print('-i, --index\t序号')
            sys.exit(0)
        if key in ['-i', '--index']:
            index = value
    print('马尔代夫V5.5 function with num: ', index ) 

# STEP 1，首先读取所需的paragraph.csv，detail.csv文件，被注释的两端分别用来检测是否运行成功

    # 查找所选index，这里利用变量pick保存（字典类型）
    tables = excel_table_byindex('主题分配.xlsx', 0, 1) 
    pick_temp = tables[int(index) - 1]
    pick = {}
    
    for key in pick_temp:
        pick[key] = str(pick_temp[key])
         
    ref_app = []
    with codecs.open('app.csv', "r") as f:
        reader = csv.reader(f)
        for row in reader:
            ref_app.append(row[0])    
        f.close()
                
#    # 测试输出，查看是否正确，讲道理不会出问题的，看雇的测试员怎么说
#    print(pick)

    # 读取本地库，row是list表，我们统计有数据的列，利用temp_num存储每行数量
    # 利用temp_val存储每行的具体值，利用ref_num,ref_val存储整个paragraph.csv文件
    ref_num = []
    ref_val = []
    with codecs.open('paragraph.csv', "r") as f:
        reader = csv.reader(f)
        for row in reader:
            temp_num = 0
            temp_val = []
            for i, val in enumerate(row):
                if i == 0:
                    pass
                elif val != '':
                    temp_num += 1
                    temp_val.append(val)
            ref_num.append(temp_num)
            ref_val.append(temp_val)    
        f.close()
        
#    # 测试是否储存成功        
#    for i, temp_num in enumerate(ref_num):
#        for j in ref_val[i]:
#            print(j)  

    # 查找详情
    detail = excel_table_byindex('详情.xlsx', 0, 0)
    bl = detail[0]['参数']
    bw = detail[1]['参数']
    h = detail[2]['参数']
    tl = detail[3]['参数']
    tw = detail[4]['参数']
    q1 = detail[7]['电器参数']
    q2 = detail[8]['电器参数']
    i1 = detail[10]['电器参数']
    i2 = detail[11]['电器参数']
    v1 = detail[13]['电器参数']
    v2 = detail[14]['电器参数']
    r = detail[16]['电器参数']
    feature = []
    for i in range(30):
        if detail[i]['特性'] == '':
            break
        else:
            temp = detail[i]['特性']
            t = temp.split('• ')
            feature.append(t[1])

# STEP 2，三段话的随机组合生成

#    # 伪随机数生成模块。如果不提供 seed，默认使用系统时间。使用相同的 seed，可以获得完全相同的随机数序列，常用于算法改进测试。
#    random.seed(10086)

    # list 格式的三段话
    pl1 = []
    pl2 = []
    pl3 = []        

    # 生成三段话
    for i, temp_num in enumerate(ref_num):
        if i < 3:
            r1 = random.randint(0, ref_num[i] - 1)
            pl1.append(ref_val[i][r1])
        elif i < 4:
            r2 = random.randint(0, ref_num[i] - 1)
            pl2.append(ref_val[i][r2])   
        else:
            r3 = random.randint(0, ref_num[i] - 1)
            pl3.append(ref_val[i][r3])
            
#    # 测试输出结果        
#    print(pl1)
#    print(pl2)
#    print(pl3)

# STEP 3，参数的替换
    r_list = range(0, len(ref_app))
    app_index = random.sample(r_list, 3)

    temp1 = pl1[0].replace('XXX', pick['型号'])
    p1 = temp1 + pl1[1] + pl1[2]
    
    temp1 = pl2[0].replace('OOO', pick['型号'])
    temp1 = temp1.replace('XXX', str(bl) + ' X ' + str(bw))
    temp1 = temp1.replace('YYY', str(tl) + ' X ' + str(tw))
    temp1 = temp1.replace('ZZZ', str(h))
    p2 = temp1
    
    temp1 = pl3[0].replace('XXX', str(q1))
    temp1 = temp1.replace('YYY', str(i1))
    temp1 = temp1.replace('ZZZ', str(v1))
    temp1 = temp1.replace('MMM', str(r))
    temp1 = temp1.replace('NNN', str(q2))
    temp1 = temp1.replace('PPP', str(i2))
    temp1 = temp1.replace('QQQ', str(v2))
    p3 = temp1
    
#    # 测试输出结果 
#    print(p1)
#    print(p2)
#    print(p3)

# STEP 4，word的保存

    # 标题由detail导入
    MyContent_type = '新产品'
    MyContent_abstract = ''  # 摘要默认为空，手写
    MyContent_factory = 'II-VI Marlow（贰陆马洛）'
    MyContent_device = '制冷片，单级半导体制冷片，Single-Stage Thermoelectric Module'
    # 型号由detail导入
    MyContent_app = '见文章内容'  # 应用选用随机生成
    MyContent_key = '热端温度，最大功率，最大电流，最大电压，交流电阻，模块高度'
    MyContent_name = '刘晨（翻译）'
    MyContent_author = '泊棠'
    # 参考链接由detail导入

    doc = DocxTemplate("templet.docx")
    my_title1 = '图1：' + pick['型号'] + '示意图'
    my_title2 = '表1：' + pick['型号'] + '电器规格表'
    context = {
        'MyTable' : [
            {'title' : '【产品】', 'type' : MyContent_type, 'abstract' : MyContent_abstract,
             'factory' : MyContent_factory, 'device' : MyContent_device, 'version' : pick['型号'], 
             'app' : MyContent_app, 'key' : MyContent_key, 'name' : MyContent_name, 
             'author' : MyContent_author, 'ref' : pick['参考链接'] }
        ],
        'MyFirstP' : p1 ,
        'MySecondP' : p2 ,
        'MyThirdP' : p3 ,
        'MyTitle1' : my_title1 ,
        'MyTitle2' : my_title2
    }
    
    doc.render(context)
    doc.save("generated_temp.docx")
    
    doc = Document("generated_temp.docx")
    
    p1 = doc.add_paragraph('')
    r1 = p1.add_run(pick['型号'] + '的主要特点：')
    for i in feature:
        if i[-1] == '.':
            temp = i[:-1]
        else:
            temp = i
        doc.add_paragraph('• ' + temp)
        
    doc.add_paragraph('')
    p2 = doc.add_paragraph('')
    r2 = p2.add_run(pick['型号'] + '的典型应用：')
    
    doc.add_paragraph(ref_app[app_index[0]]) 
    doc.add_paragraph(ref_app[app_index[1]]) 
    doc.add_paragraph(ref_app[app_index[2]]) 
    
    file_name = 'result_temp.docx'
    
    doc.save(file_name)  
