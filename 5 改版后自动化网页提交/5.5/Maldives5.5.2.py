# -*- coding: utf-8 -*-
"""
Created on Fri Dec 29 18:42:11 2017

@author: liurole

马尔代夫V5.5.2

实现了文章由word的提交

配置文件：
result_temp.docx 修改后的文件，图片不用粘贴

如果提示缺少库就通过pip install安装
另外需要注意的是需要安装Chrome以及驱动
下载chromedrive驱动 
使用Selenium需要选择一个调用的浏览器并下载好对应的驱动，本文使用chrome浏览器，当然也可以用FireFox等
http://www.seleniumhq.org/download/ 找到Google Chrome Driver链接
对应驱动放在python目录下面的scripts目录中，例如C:\ProgramData\Anaconda3\envs\python35\Scripts
注：如果是macOS，下载对应版本驱动解压到环境变量包含的路径即可，比如/usr/local/bin

"""

# 这些都是配置库，不用管
import time
from docx import Document
from selenium import webdriver

if __name__ == '__main__':

# STEP 1，读入文档数据

    #打开文档
    document = Document('result_temp.docx')
    #读取每段资料
    para = [ paragraph.text for paragraph in document.paragraphs];
    #输出并观察结果，也可以通过其他手段处理文本即可

    #读取表格材料，并输出结果
    store = []
    for table in document.tables:
        for row in table.rows:
            store_temp = []
            for cell in row.cells:
                store_temp.append(cell.text)
            store.append(store_temp)

    MyContent_title = store[0][1]
    MyContent_type = store[1][1]
    MyContent_abstract = store[2][1]
    MyContent_factory = store[3][1]
    MyContent_device = store[4][1]
    MyContent_version = store[5][1]
    
    
    MyContent_app = store[6][1]
    MyContent_key = store[7][1]
    MyContent_name = store[8][1]
    MyContent_author = store[9][1]
    MyContent_web = store[10][1]
    
    file_name = '刘晨+' + MyContent_version + '.docx'
    document.save(file_name)  

# STEP 2，读入文档数据 发布文章
    
    # 打开chorme
    bolgurl = 'https://cms.sekorm.com/snps/'
    browser = webdriver.Chrome()
    browser.get(bolgurl)
    
    # 输入验证码
    name = browser.find_element_by_id('code')
    name.send_keys('zneWR56ylKQ1Pvtn')
    
    # 点击确定，加载完成后等待4秒（网页设定为3s）
    login_button = browser.find_element_by_xpath('/html/body/div[2]/div/div[2]/div/div[1]/div[1]/form/div[3]/div/div/div/div/button')
    login_button.click()
    time.sleep(5)
    # browser.implicitly_wait(4)
    
    # 标题输入
    name = browser.find_element_by_id('title')
    name.send_keys(MyContent_title)
    
    # 新产品
    if MyContent_type == '新产品':
        browser.find_element_by_name('acticleType').click()
    
    # 摘要
    name = browser.find_element_by_xpath('//*[@id="root"]/div/div/div[2]/div[1]/div[2]/form/div[4]/div[1]/div/div/div[2]/div/textarea')
    name.send_keys(MyContent_abstract)
    
    # 厂牌
    if MyContent_factory == 'II-VI Marlow（贰陆马洛）':
        browser.find_element_by_id('brandCode').click()
        browser.find_element_by_id('brandCodetreeLeft_58_a').click()
        browser.find_element_by_id('brandCodeconfirmBtn').click()
    
    # 器件名称
    name = browser.find_element_by_id('goodKeyword')
    name.send_keys(MyContent_device)
    
    # 型号
    name = browser.find_element_by_id('pnKeyword')
    name.send_keys(MyContent_version)
    
    # 市场应用
    name = browser.find_element_by_id('elecKeyword')
    MyContent_app = ''
    for i in range(0, para.__len__())[::-1]:
        if i == len(para) - 1:
            MyContent_app += para[i]
        elif '的典型应用' in para[i]:
            break
        else:
            MyContent_app += '，' + para[i]
    name.send_keys(MyContent_app)
    
    # 关键词
    name = browser.find_element_by_id('keyword')
    name.send_keys(MyContent_key)
    
    # 作者姓名
    name = browser.find_element_by_id('writer')
    name.send_keys(MyContent_name)
    
    # 笔名
    name = browser.find_element_by_id('author')
    name.send_keys(MyContent_author)
    
    # 参考链接
    name = browser.find_element_by_id('referLink')
    name.send_keys(MyContent_web)
    
    
    # 正文输入
    browser.switch_to.frame('ueditor_0')  # 注意，这种editor一定有frame，一定要切frame
    MyContent_content = ''
    bapp = 0
    for i, val in enumerate(para):
        if val == '':
            MyContent_content += '\n'
        else:
            if bapp:
                MyContent_content += '• ' + val + '\n'
            else:
                MyContent_content += val + '\n'
        if '的典型应用' in val:
            bapp = not bapp
    MyContent_content += '\n世强元件电商版权所有，转载请注明来源及链接。'
    browser.find_element_by_tag_name('body').send_keys(MyContent_content)
    browser.switch_to.default_content()
        
    