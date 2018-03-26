# -*- coding: utf-8 -*-
"""
Created on Wed Feb 14 16:44:45 2018

@author: Se

下载chromedrive驱动 
使用Selenium需要选择一个调用的浏览器并下载好对应的驱动，本文使用chrome浏览器，当然也可以用FireFox等
http://www.seleniumhq.org/download/ 找到Google Chrome Driver链接
对应驱动放在python目录下面的scripts目录中，例如C:\ProgramData\Anaconda3\envs\python35\Scripts

记得改成自己的账户密码！！！

实现了文章的自动化提交，文章格式参加模板
"""
import sys
import getopt
import time
from selenium import webdriver
from docx import Document

# run Maldives6.1.0.py -i 32864115 -p 1
if __name__ == '__main__':
    opts, args = getopt.getopt(sys.argv[1:], 'hi:p:', [ 'help', 'index = ', 'passed = ' ])
    
    index = str(32864115)
    passed = str(1)
    
    # 入口函数，不明白怎么调用参数的可以看下
    for key, value in opts:
        if key in ['-h', '--help']:
            print('参数定义：')
            print('-h, --help\t显示帮助')
            print('-i, --index\t序号')
            print('-p, --passed\t主题是否通过')
            sys.exit(0)
        if key in ['-i', '--index']:
            index = str(value)
        if key in ['-p', '--passed']:
            passed = str(value)

    print('一定记得修改为自己的账号密码！\t主题序号：' + index + '\t是否通过:' + passed)
    
    # 打开chorme
    url = 'https://cms.sekorm.com/content/login'
    browser = webdriver.Chrome()
    browser.get(url)
    time.sleep(1)
    
    # 输入验证码
    name = browser.find_element_by_id('mobile')
    name.send_keys('15244608508')
    password = browser.find_element_by_id('password')
    password.send_keys('liu875288')
    button = browser.find_element_by_xpath('//*[@id="root"]/div/div/div[2]/div/div[1]/div/div/div/div[2]/form/div[3]/div/button')
    button.click()
    
    # 切换到提文章页面
    url = 'https://cms.sekorm.com/content/nps/pendingSubject/edit?id=' + index + '&status=' + passed + '&ecnewStatus=null'
    browser.get(url)
    time.sleep(3)
    
    #打开文档
    document = Document('模板.docx')
    #读取每段资料
    para = [ paragraph.text for paragraph in document.paragraphs];
    #输出并观察结果，也可以通过其他手段处理文本即可

    #读取表格材料，并输出结果
    store = []
    for table in document.tables:
        for row in table.rows:
            store_temp = []
            for cell in row.cells:
                store_temp.append(cell.text)
            store.append(store_temp)

    MyContent_title = store[0][1]
    MyContent_factory = store[1][1]
    MyContent_device = store[2][1]
    MyContent_version = store[3][1]
    MyContent_app = store[4][1]
    MyContent_key = store[5][1]
    MyContent_abstract = store[6][1]

    
    # 标题输入
    name = browser.find_element_by_id('title')
    name.send_keys(MyContent_title)
    
#    # 厂牌
#    if MyContent_factory == 'Kyocera(京瓷)':
#        browser.find_element_by_id('brand').click()
#        time.sleep(2)
#        browser.find_element_by_id('brandtreeLeft_47_span').click()
#        browser.find_element_by_id('brandconfirmBtn').click()   
#    elif MyContent_factory == 'shindengen（新电元）':
#        browser.find_element_by_id('brand').click()
#        time.sleep(2)
#        browser.find_element_by_id('brandtreeLeft_24_span').click()
#        browser.find_element_by_id('brandconfirmBtn').click()  
#    
#    # 器件名称
#    name = browser.find_element_by_id('goods')
#    name.send_keys(MyContent_device)
#    
#    # 型号
#    name = browser.find_element_by_id('pn')
#    name.send_keys(MyContent_version)
    
    
    # 应用
    name = browser.find_element_by_id('elec')
    name.send_keys(MyContent_app)
    
    # 关键词
    name = browser.find_element_by_id('keyword')
    name.send_keys(MyContent_key)
    
    # 摘要
    name = browser.find_element_by_xpath('//*[@id="root"]/div/div[2]/div[2]/div[2]/div/div/div/div/div/div/div[2]/div/form/div[2]/div[7]/div/div/div/div[2]/div/span/textarea')
    name.send_keys(MyContent_abstract)

    file_name = MyContent_title + '.docx'
    document.save(file_name)  

    # 正文输入
    browser.switch_to.frame('ueditor_0')  # 注意，这种editor一定有frame，一定要切frame
    MyContent_content = ''
    bapp = 0
    for i, val in enumerate(para):
        if val == '':
            MyContent_content += '\n'
        else:
            MyContent_content += val + '\n'
                
    browser.find_element_by_tag_name('body').send_keys(MyContent_content)
    browser.switch_to.default_content()
    
#    # 注释掉此处可以自动提交草稿，保险起见，你还是手动检查一下吧
#    # 保存草稿
#    browser.find_element_by_id('saveAsBackup').click()  