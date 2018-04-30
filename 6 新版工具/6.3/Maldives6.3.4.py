# -*- coding: utf-8 -*-
"""
Created on Wed Feb 14 22:03:08 2018

@author: Se

下载chromedrive驱动 
使用Selenium需要选择一个调用的浏览器并下载好对应的驱动，本文使用chrome浏览器，当然也可以用FireFox等
http://www.seleniumhq.org/download/ 找到Google Chrome Driver链接
对应驱动放在python目录下面的scripts目录中，例如C:\ProgramData\Anaconda3\envs\python35\Scripts
"""
import sys
import getopt
import os
import time
from selenium import webdriver
from urllib.parse import urlencode
import requests
import re
import urllib
# 可直接通过pip install 安装
from docx import Document
from docxtpl import DocxTemplate
from bs4 import BeautifulSoup

# 利用request，导入cookies，header进行关键词网页搜索，选择第二栏
def check(url):
    try:
        response = requests.get(url)
        response.encoding="utf-8"
        if response.status_code == 200:
            return response.text
        return None
    except ConnectionError:
        print('Error occurred')
        return None

# 获取页面信息
def search_url(url):
    html = check(url)
    soup = BeautifulSoup(html, 'lxml')
    result = soup.select('#activity-name')
    title = result[0].get_text()
    title = title.strip()
    
    # 日期
    result = soup.select('#post-date')
    date = result[0].get_text()
    
    # 信息获取
    MyContent_url = url
    MyContent_title = title
    MyContent_factory = 'Silicon Labs'
    MyContent_device = ''
    MyContent_version = ''
    MyContent_app = ''
    MyContent_key = ''
    MyContent_abstract = ''
    MyContent_date = date
    MyContent_type = '新产品'
    doc = DocxTemplate("templet.docx")
    context = {
        'MyTable' : [
            {'ref' : MyContent_url, 'title' : MyContent_title, 'factory' : MyContent_factory, 'device' : MyContent_device,
             'version' : MyContent_version, 'app' : MyContent_app, 'key' : MyContent_key, 
             'abstract' : MyContent_abstract, 'date' : MyContent_date, 'type' : MyContent_type}
        ]
    }    
    doc.render(context)
    doc.save("generated_temp.docx")
    
    # 保存图片
    result = soup.select('#js_content img')
    index = 1
    for i in result:
        link = i.get('data-src')
        filename = path + '\\' + str(index) + '.jpg'
        urllib.request.urlretrieve(link, filename)        
    
    # 查找内容
    result = soup.select('#js_content p')
    
    return result, title
    
# 提交单个系列产品，仅适用于新产品
def submit(browser, surl, result, title):    
    browser.get(surl)
    time.sleep(2)    
    
    # 内容
    text = ''
    doc = Document("generated_temp.docx")       
    for i in result:
        temp = i.get_text()
        if temp == '':
            text += '\n'
            doc.add_paragraph('')            
        else:
            text += temp + '\n'
            doc.add_paragraph(temp)

    file_name = title + '.docx'
    doc.save(file_name)  
    
    # 正文输入
    browser.switch_to.frame('ueditor_0')  # 注意，这种editor一定有frame，一定要切frame
    browser.find_element_by_tag_name('body').send_keys(text)
    browser.switch_to.default_content()
              
    print('------Done!')

# 帮助run Maldives6.3.0.py --help
# run Maldives6.3.1.py -i 95523562 -p 1 -u http://www.eeworld.com.cn/qrs/article_2018042547518.html
if __name__ == '__main__':
    opts, args = getopt.getopt(sys.argv[1:], 'hi:p:u:', [ 'help', 'index = ', 'passed = ', 'url=' ])
    
    url = ''
    index = ''
    passed = ''
    path = 'D:\\sekorm'
    isExists = os.path.exists(path)
    if not isExists:
        os.makedirs(path)
    
    # 入口函数，不明白怎么调用参数的可以看下
    for key, value in opts:
        if key in ['-h', '--help']:
            print('马尔代夫V6.3')
            print('参数定义：')
            print('-h, --help\t显示帮助')
            print('-u, --url\t参考链接')
            sys.exit(0)
        if key in ['-u', '--url']:
            url = value
        if key in ['-i', '--index']:
            index = str(value)
        if key in ['-p', '--passed']:
            passed = str(value)
    
    if index == '' or  passed == '' or  url == '':
        exit()
        
    print('马尔代夫V6.3 function with: ' + '\t' + index) 
    
    result, title = search_url(url)
    
    # 打开chorme
    surl = 'https://cms.sekorm.com/content/login'
    browser = webdriver.Chrome()
    browser.get(surl)
    
    # 输入验证码
    name = browser.find_element_by_id('mobile')
    name.send_keys('15244608508')
    password = browser.find_element_by_id('password')
    password.send_keys('liu875288')
    button = browser.find_element_by_xpath('//*[@id="root"]/div/div/div[2]/div/div[1]/div/div/div/div[2]/form/div[3]/div/button')
    button.click()    
    
    # 切换到提文章页面
    surl = 'https://cms.sekorm.com/content/nps/pendingSubject/edit?id=' + index + '&status=' + passed + '&ecnewStatus=null'

    submit(browser, surl, result, title)
    
    