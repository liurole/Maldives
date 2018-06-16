# -*- coding: utf-8 -*-
"""
Created on Fri Dec 29 18:42:11 2017

@author: liurole

配置文件：
result_temp.docx 修改后的文件，图片不用粘贴

如果提示缺少库就通过pip install安装
另外需要注意的是需要安装Chrome以及驱动
下载chromedrive驱动 
使用Selenium需要选择一个调用的浏览器并下载好对应的驱动，本文使用chrome浏览器，当然也可以用FireFox等
http://www.seleniumhq.org/download/ 找到Google Chrome Driver链接
对应驱动放在python目录下面的scripts目录中，例如C:\ProgramData\Anaconda3\envs\python35\Scripts
注：如果是macOS，下载对应版本驱动解压到环境变量包含的路径即可，比如/usr/local/bin

"""

# 这些都是配置库，不用管
import sys
import getopt
import time
from docx import Document
from selenium import webdriver

#run Maldives6.4.3.py -i 95000124 -p 1
if __name__ == '__main__':
    opts, args = getopt.getopt(sys.argv[1:], 'hi:p:', [ 'help', 'index = ', 'passed = ' ])
    
    index = ''
    passed = ''
    
    # 入口函数，不明白怎么调用参数的可以看下
    for key, value in opts:
        if key in ['-h', '--help']:
            print('马尔代夫V6.4')
            print('参数定义：')
            print('-h, --help\t显示帮助')
            print('-u, --url\t参考链接')
            sys.exit(0)
        if key in ['-i', '--index']:
            index = str(value)
        if key in ['-p', '--passed']:
            passed = str(value)
    
    if index == '' or  passed == '':
        sys.exit()
        
    print('马尔代夫V6.4 function with: ' + '\t' + index)     
    
# STEP 1，读入文档数据

    #打开文档
    document = Document('result_temp.docx')
    #读取每段资料
    para = [ paragraph.text for paragraph in document.paragraphs];
    #输出并观察结果，也可以通过其他手段处理文本即可

    #读取表格材料，并输出结果
    store = []
    for table in document.tables:
        for row in table.rows:
            store_temp = []
            for cell in row.cells:
                store_temp.append(cell.text)
            store.append(store_temp)

    MyContent_title = store[0][1]
    MyContent_app = store[1][1]
    MyContent_key = store[2][1]
    MyContent_abstract = store[3][1]
    
    file_name = './已撰写/' + MyContent_title + '.docx'
    document.save(file_name) 

# STEP 2，发布文章
    # 打开chorme
    surl = 'https://cms.sekorm.com/content/login'
    browser = webdriver.Chrome()
    browser.get(surl)
    
    # 输入验证码
    name = browser.find_element_by_id('mobile')
    name.send_keys('15244608508')
    password = browser.find_element_by_id('password')
    password.send_keys('liu875288')
    button = browser.find_element_by_xpath('//*[@id="root"]/div/div/div[2]/div/div[1]/div/div/div/div[2]/form/div[3]/div/button')
    button.click()    
    
    # 切换到提文章页面
    surl = 'https://cms.sekorm.com/content/nps/pendingSubject/edit?id=' + index + '&status=' + passed + '&ecnewStatus=null'

    browser.get(surl)
    time.sleep(2)  
    
    # 标题输入
    name = browser.find_element_by_id('title')
    name.send_keys(MyContent_title)

    # 应用输入
    name = browser.find_element_by_id('elec')
    name.send_keys(MyContent_app)

    # 关键词
    name = browser.find_element_by_id('keyword')
    name.send_keys(MyContent_key)

    # 摘要
    name = browser.find_element_by_xpath('//*[@id="root"]/div/div[2]/div[2]/div[2]/div/div/div/div/div/div/div[2]/div/form/div[2]/div[7]/div/div/div/div[2]/div/span/textarea')
    name.send_keys(MyContent_abstract)

    # 正文输入
    browser.switch_to.frame('ueditor_0')  # 注意，这种editor一定有frame，一定要切frame
    MyContent_content = ''
    bapp = 0
    for i, val in enumerate(para):
        if val == '':
            MyContent_content += '\n'
        else:
            MyContent_content += val + '\n'
                
    browser.find_element_by_tag_name('body').send_keys(MyContent_content)
    browser.switch_to.default_content()
