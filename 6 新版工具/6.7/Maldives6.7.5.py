# -*- coding: utf-8 -*-
"""
Created on Fri Dec 29 08:03:29 2017

@author: liurole


针对于京瓷肖特基二极管
尝试系列自动化

实现了文章的自动化网页生成同时word本地备份，防止撕逼

需指定生成index（-i []）

配置文件：
templet.docx 模板文件
paragraph.csv 库函数文件，此处更新库语句
detail.csv 特性文件，此文件为所有产品的信息集合
app.csv 应用文件

定义一下各文件格式：
templet.docx 变数不大，这个到时候我来改就行
paragraph.csv 这个是变数最大的文件，考虑到语句数量可能不同，这里每一句用的是行排列
detail.csv 这里加一个表头（第一行），方便调用，同时添加一列序号，方便每次生成

"""

# 这些都是配置库，不用管
import sys
import getopt
import random
import xlrd  

# 可直接通过pip install 安装
from docx import Document
from docxtpl import DocxTemplate

def excel_table_byindex(colnameindex = 0, by_index = 0):  
    app_para = 'paragraph.xlsx'
    app_file = 'app.xlsx'
    detail_file = '京瓷肖特基汇总.xlsx'
    
    # 读取段落
    data = xlrd.open_workbook(app_para)  
    table = data.sheets()[0] 
    nrows = table.nrows #行数
    para = []
    for rownum  in range(nrows):
        temp = []
        row = table.row_values(rownum)
        for i in row:
            if i != '':
                temp.append(i)
        para.append(temp)
   
    # 读取应用
    data = xlrd.open_workbook(app_file)  
    table = data.sheets()[0] 
    nrows = table.nrows #行数 
    app = []
    for rownum  in range(nrows):
        row = table.row_values(rownum)
        app.append(row[0])
    
    # 读取详情
    data = xlrd.open_workbook(detail_file)  
    table = data.sheets()[by_index]  
    nrows = table.nrows #行数   
    colnames = table.row_values(colnameindex) #某一行数据  
    detail = []  
    for rownum in range(1,nrows):  
        row = table.row_values(rownum)#以列表格式输出  
        if row:  
            temp = {}  
            for i in range(len(colnames)):  
                temp[colnames[i]] = row[i]  
            detail.append(temp)#向列表中插入字典类型的数据  
    print('Read Done!')
    
    return app, para, detail  

# 直接修改程序吧
if __name__ == '__main__':
    opts, args = getopt.getopt(sys.argv[1:], 'hi:', [ 'help', 'index=' ])
    
    index = [174,175]
    #index = 133
    
    # 入口函数，不明白怎么调用参数的可以看下
    for key, value in opts:
        if key in ['-h', '--help']:
            print('马尔代夫V6.7')
            print('参数定义：')
            print('-h, --help\t显示帮助')
            print('-i, --index\t序号')
            sys.exit(0)
        if key in ['-i', '--index']:
            index = value
            
    if index ==None:
        sys.exit()        
    print('马尔代夫V6.7 function with num: ', index ) 

    # 读取记录
    app, para, detail  = excel_table_byindex(0, 0)
    
    # 获取所选的系列产品的所有信息
    pick = {}
    if type(index) == int:
        pick = detail[index - 1]
        print('单独型号！')
        for key in pick:
            temp = str(pick[key])  
            if temp[-2:] == '.0':
                temp = temp[:-2]
            pick[key] = str(temp)
    else:
        print('系列型号！')
        pick_temp = detail[index[0] - 1]
        for key in pick_temp:
            temp = str(pick_temp[key])  
            if temp[-2:] == '.0':
                temp = temp[:-2]
            pick[key] = str(temp)
        
        for i in range(1, len(index)):
            pick_temp = detail[index[i] - 1]
            for key in pick_temp:
                temp = str(pick_temp[key])    
                if temp[-2:] == '.0':
                    temp = temp[:-2]
                if temp != pick[key]:
                    pick[key] += '/' + temp
    
    # 随机选取三个应用
    r_list = range(0, len(app))
    app_index = random.sample(r_list, 2)
    
    pl1 = ''
    pl2 = ''
    pl3 = '' 
    # 生成第一段
    r = random.randint(0, len(para[0]) - 1)
    temp1 = para[0][r]
    temp1 = temp1.replace('XXX', str(pick['name']))  
    r = random.randint(0, len(para[1]) - 1)
    temp2 = para[1][r]
    r = random.randint(0, len(para[2]) - 1)
    temp3 = para[2][r]
    r = random.randint(0, len(para[3]) - 1)
    temp4 = para[3][r]
    r = random.randint(0, len(para[4]) - 1)
    temp5 = para[4][r]
    pl1 = temp1 + temp2 + temp3 + temp4 + temp5
    
    # 生成第二段
    r = random.randint(0, len(para[5]) - 1)
    temp1 = para[5][r]
    temp1 = temp1.replace('XXX', str(pick['name']))  
    temp1 = temp1.replace('YYY', str(pick['VRRM']))
    r = random.randint(0, len(para[6]) - 1)
    temp2 = para[6][r]
    temp2 = temp2.replace('XXX', str(pick['IO']))
    temp2 = temp2.replace('YYY', str(pick['IFSM']))
    r = random.randint(0, len(para[7]) - 1)
    temp3 = para[7][r]
    temp3 = temp3.replace('XXX', str(pick['name'])) 
    temp3 = temp3.replace('YYY', str(pick['VRRM'])) 
    temp3 = temp3.replace('ZZZ', str(pick['IR'])) 
    pl2 = temp1 + temp2 + temp3
    
    # 生成第三段
    r = random.randint(0, len(para[8]) - 1)
    temp1 = para[8][r]
    temp1 = temp1.replace('XXX', str(pick['name']))  
    temp1 = temp1.replace('YYY', str(pick['VFM']))
    temp2 = para[9][0]
    temp = pick['Tstg']
    temp = temp.replace(' ～ ', '℃至')
    temp2 = temp2.replace('XXX', temp)
    r = random.randint(0, len(para[10]) - 1)
    temp3 = para[10][r]
    pl3 = temp1 + temp2 + temp3
    
    # 写入应用
    doc = DocxTemplate("templet.docx")
    context = {
        'MyTable' : [
            {'title' : '【产品】', 'app' : '二次侧整流, DC/DC转换，逆流保护', 'key' : '最大反向电压，正向峰值浪涌电流, 反向漏电流, 平均正向整流电流, 正向峰值电压', 'abstract' : ''}
        ]
    }   
    doc.render(context)
    doc.save("generated_temp.docx")
                
    # 写入word
    doc = Document("generated_temp.docx")
    doc.add_paragraph(pl1)
    doc.add_paragraph('')
    doc.add_paragraph(pl2)
    doc.add_paragraph('')
    doc.add_paragraph('')
    if '/' in pick['name']:
        temp = pick['name'].split('/')
        doc.add_paragraph('图1  ' + temp[0] + '的封装示意图')
    else:    
        doc.add_paragraph('图1  ' + pick['name'] + '的封装示意图')
    doc.add_paragraph('')
    doc.add_paragraph(pl3)
    doc.add_paragraph('')
    doc.add_paragraph('')
    if '/' in pick['name']:
        temp = pick['name'].split('/')
        doc.add_paragraph('图2  ' + temp[0] + '的瞬时正向电压特性曲线')
    else:
        doc.add_paragraph('图2  ' + pick['name'] + '的瞬时正向电压特性曲线')
    doc.add_paragraph('')  
    
    doc.add_paragraph(pick['name'] + '的主要特点：')
    temp = '最大反向电压为XXXV，结温为25℃的情况下，反向漏电流为YYYmA'
    temp = temp.replace('XXX', str(pick['VRRM']))
    temp = temp.replace('YYY', str(pick['IR']))
    doc.add_paragraph('• ' + temp) 
    temp = '平均正向整流电流为XXXA, 正向峰值浪涌电流为YYYA,'
    temp = temp.replace('XXX', str(pick['IO']))
    temp = temp.replace('YYY', str(pick['IFSM']))
    doc.add_paragraph('• ' + temp) 
    temp = '正向峰值电压为XXXV'
    temp = temp.replace('XXX', str(pick['VFM']))
    doc.add_paragraph('• ' + temp) 
    doc.add_paragraph('• 适用于高频应用') 
    doc.add_paragraph('• 环氧树脂封装，阻燃等级符合UL94V-0标准') 
    doc.add_paragraph('• 符合RoHS标准') 
    

    doc.add_paragraph('')
    doc.add_paragraph(pick['name'] + '的典型应用：')
    doc.add_paragraph('• 二次侧整流') 
    doc.add_paragraph('• DC/DC转换') 
    doc.add_paragraph('• 逆流保护') 
    for i in app_index:
        doc.add_paragraph('• ' + app[i])
    doc.add_paragraph('')
    
    file_name = 'result_temp.docx'
    
    doc.save(file_name)  
