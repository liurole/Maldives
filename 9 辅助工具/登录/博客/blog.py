# -*- coding: utf-8 -*-
"""
Created on Wed Feb 14 22:03:08 2018

@author: Se

下载chromedrive驱动 
使用Selenium需要选择一个调用的浏览器并下载好对应的驱动，本文使用chrome浏览器，当然也可以用FireFox等
http://www.seleniumhq.org/download/ 找到Google Chrome Driver链接
对应驱动放在python目录下面的scripts目录中，例如C:\ProgramData\Anaconda3\envs\python35\Scripts
"""
import requests
import re
import urllib
import hashlib
import random
# 可直接通过pip install 安装
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docxtpl import DocxTemplate
from bs4 import BeautifulSoup

def baidu_fanyi(q: str, fro: str = 'auto', to: str = 'zh', timeout: int = 5):
    """
    调用百度翻译API实现在线翻译
    :param q:请求翻译query
    :param fro:翻译源语言，默认自动识别
    :param to:译文语言，默认zh
    :param timeout:设置超时时间，默认5秒
    :return:正常返回json格式数据，否则返回None，结果保存在['trans_result'][0]['dst']中
    """
    appid = '20171203000101916'  # fill in your app_id
    salt = str(random.random())
    key = '8W9VbfBIzhNOhYT729WW'  # fill in your key
    text = appid + q + salt + key
    md5 = hashlib.md5()
    md5.update(text.encode('UTF-8'))
    sign = md5.hexdigest()
    url = 'https://fanyi-api.baidu.com/api/trans/vip/translate'
    argv = {
        'appid': appid,
        'salt': salt,
        'key': key,
        'sign': sign,
        'q': q,
        'from': fro,
        'to': to
    }
    try:
        r = requests.get(url, params=argv, timeout=timeout)
    except:
        return None
    return r.json()

def youdao_fanyi(q: str, fro: str = 'auto', to: str = 'zh', timeout: int = 5):
    """
    调用有道翻译API实现在线翻译
    :param q:请求翻译query
    :param fro:翻译源语言，默认自动识别
    :param to:译文语言，默认zh
    :param timeout:设置超时时间，默认5秒
    :return:正常返回json格式数据，否则返回None，结果保存在['translation']中
    """
    appid = '68a68ddaf01f7b2f'  # fill in your app_id
    salt = str(random.random())
    key = 'AQwzCoZjhpMNVZApqUb9PWEghKv9eBLD'  # fill in your key
    text = appid + q + salt + key
    md5 = hashlib.md5()
    md5.update(text.encode('UTF-8'))
    sign = md5.hexdigest()
    url = 'https://openapi.youdao.com/api/'

    argv = {
        'appKey': appid,
        'q': q,
        'from': fro,
        'to': to,
        'salt': salt,
        'sign': sign
    }
    try:
        r = requests.get(url, params=argv, timeout=timeout)
    except:
        return None
    return r.json()

# 利用request，导入cookies，header进行关键词网页搜索，选择第二栏
def check(url):
    headers = {
    'Accept': 'text/html, application/xhtml+xml, image/jxr, */*',
    'Accept-Encoding': 'gzip, deflate',
    'Accept-Language': 'zh-CN',
    'Connection': 'keep-alive',
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/51.0.2704.79 Safari/537.36 Edge/14.14393'
    }
    try:
        response = requests.get(url, headers = headers)
        if response.status_code == 200:
            return response.text
        return None
    except ConnectionError:
        print('Error occurred')
        return None

# 获取所有搜索结果链接
def get_all_urls(html):
    names = []
    urls = []
    soup = BeautifulSoup(html, 'lxml')
    result = soup.select('#archive .strong a')
    if len(result) == 0:
        return names, urls
    else:
        for i in result:
            temp = i.get('href')
            urls.append('https://www.planetanalog.com' + temp)
            temp = i.get_text()
            names.append(temp)
        return names, urls

# 保存所有页面
def search_all_urls(url):
    response = check(url)
    result =re.search('<div class="grayshowlinks bigsmall" style="margin-bottom: 14px;">.*<div class="divsplitter" style="height: 1px;"></div>', response, re.S)
    result = result.group()
    # 得到标题，段落，图片
    soup = BeautifulSoup(response, 'lxml')
    title = soup.select('#thedoctop .biggest')
    temp = title[0].get_text()
    temp_ch = baidu_fanyi(temp)
    temp_ch = temp_ch['trans_result'][0]['dst']
    doc = DocxTemplate('templet.docx')
    context = {
        'MyTitle1' : temp ,
        'MyTitle2' : temp_ch
    }
    file_name = re.sub('[^A-Za-z\s]','',temp) + '.docx'
    doc.render(context)
    doc.save("generated_temp.docx")
    
    paras = re.findall('<p.*?>(.*?)</p>', result)
    pics = re.findall('<img class="docimage" src="(.*?)" alt', result)
    
    doc = Document("generated_temp.docx")
    index = 0
    for para in paras:
        if '<b>' in para:
            temp = pics[index]
            temp_split = temp.split('.')
            file = 'temp.' + temp_split[-1]
            urllib.request.urlretrieve(temp, file)
            index += 1
            doc.add_picture(file)
            
            temp = re.search('>(.*)<', para)
            if temp is not None:
                para = temp.group(1)                
            temp_ch = baidu_fanyi(para)
            temp_ch = temp_ch['trans_result'][0]['dst']
            picname = '图' + str(index) + ': '
            p1 = doc.add_paragraph(picname + para)
            p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER 
            p2 = doc.add_paragraph(picname + temp_ch)   
            p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER 
            doc.add_paragraph('')
            
        else:
            para = re.sub('<.*?>', '', para)              
            temp_ch = baidu_fanyi(para)
            temp_ch = temp_ch['trans_result'][0]['dst']
            doc.add_paragraph(para)
            doc.add_paragraph(temp_ch)    
            doc.add_paragraph('')
               
    doc.save(file_name)              
    print(file_name + '------Done!')

if __name__ == '__main__':
    
    url = 'https://www.planetanalog.com/archives.asp?blogs=yes'
    html = check(url)
    names, urls = get_all_urls(html)
    
    for i in urls:
        search_all_urls(i)
    
    