# -*- coding: utf-8 -*-
"""
Created on Tue Dec 19 00:29:40 2017

@author: Se

马尔代夫V2.0

实现了程序化的word文档写入

依赖项：
pip install docx
pip install docxtpl

"""

from docx import Document
from docxtpl import DocxTemplate

MyContent_title = '【产品】高增益，面向2110MHz-2170MHz频段的150W高品质射频放大器，适合多标准蜂窝网功放应用'
MyContent_type = '新产品'
MyContent_factory = 'Infineon'
MyContent_version = 'PXFC211507SC'
MyContent_abstract = 'PXFC211507SC是一款高增益，带宽输入输出内匹配的面向2110MHz-2170MHz频段的功率放大器，低热阻，轻松应对恶劣环境'
MyContent_key = '射频放大器，功率放大器，LDMOS'
MyContent_app = '工业，通信，消费'
MyContent_category = '改写'
MyContent_ref = 'https://www.infineon.com/dgdl/Infineon-PXFC211507SC-DS-v02_00-EN.pdf?fileId=5546d4624c330ffd014c4830128b1259'
MyContent_author = '牛逼哥'

MyContent_firstP = '针对多标准蜂窝网络功率放大应用，德国Infineon（英飞凌）公司推出了一款150W级RF FET，型号为PXFC211507SC。\
良好的散热不仅降低了系统热设计成本，同时提升了系统在高温环境的可靠性。'
MyContent_secondP = 'PXFC211507SC可广泛应用于2110MHz-2170MHz频段的功率放大器中，具有20.4dB的增益，\
无需外围匹配电路即可实现宽带的输入输出匹配。其基板面积仅为26.3mm × 9.78mm（见图3），可减小PCB板的面积，易于实现产品的小型化。'
MyContent_thirdP = '2170MHz，28V下，PXFC211507SC的单载波WCDMA性能如下，输出功率为32W，增益可达20dB，效率为32%。'

arr1 = []
arr1.append('宽带内输入输出匹配')
arr1.append('2170MHz，28V下，其CW性能如下，P1dB输出功率为150W，增益可达19dB，效率高达56%')
arr1.append('2170MHz，28V下，其单载波WCDMA性能如下，输出功率为32W，增益可达20dB，效率32%')

arr2 = []
arr2.append('移动通信设备')
arr2.append('多标准蜂窝功率放大器')
arr2.append('蜂窝基础设施')

doc = DocxTemplate("templet.docx")
context = {
    'MyTable' : [
        {'title' : MyContent_title, 'type' : MyContent_type, 'factory' : MyContent_factory, 'version' : MyContent_version, 
         'abstract' : MyContent_abstract, 'key' : MyContent_key, 'app' : MyContent_app, 'category' : MyContent_category, 
         'ref' : MyContent_ref, 'author' : MyContent_author }
    ],
    'MyFirstP' : MyContent_firstP ,
    'MySecondP' : MyContent_secondP ,
    'MyThirdP' : MyContent_thirdP
}


doc.render(context)
doc.save("generated_temp.docx")


# 源文件 test.docx
doc = Document("generated_temp.docx")
#doc = Document("牛逼哥+PXFC211507SC.docx")
#styles = doc.styles
#
#for style in styles:
#    print(style.name)

p1 = doc.add_paragraph('')
r1 = p1.add_run('YYY的主要特点：')


for i in arr1:
    #增加无序列表
    paragraph = doc.add_paragraph(i)  
    paragraph_format = paragraph.paragraph_format  
    #paragraph_format.left_indent = Inches(0.291)  

doc.add_paragraph('')
  
p2 = doc.add_paragraph('')
r2 = p2.add_run('YYY的典型应用：')


for i in arr2:
    #增加无序列表
    paragraph = doc.add_paragraph(i)  
    paragraph_format = paragraph.paragraph_format  
    #paragraph_format.left_indent = Inches(0.291)  


doc.save("Maldives.docx")