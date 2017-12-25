# -*- coding: utf-8 -*-
"""
Created on Sat Dec  9 21:28:47 2017

@author: Se

马尔代夫V4.0

实现了文章的自动化生成

需指定生成index（-i xxx）

配置文件：
templet.docx 模板文件
paragraph.csv 库函数文件，此处更新库语句
detail.csv 特性文件，此文件为所有产品的信息集合

"""

import sys
import getopt
import csv
import random
from docx import Document
from docxtpl import DocxTemplate

# 帮助run Maldives4.0.py --help
if __name__ == '__main__':
    opts, args = getopt.getopt(sys.argv[1:], 'hi:', [ 'help', 'index=' ])
    
    index = 1
    r_num = []
    r_val = []
    
    err = []
    
    # 入口函数
    for key, value in opts:
        if key in ['-h', '--help']:
            print('马尔代夫V4.0')
            print('参数定义：')
            print('-h, --help\t显示帮助')
            print('-i, --index\t序号')
            sys.exit(0)
        if key in ['-i', '--index']:
            index = value
    print('马尔代夫V4.0 function with num: ', index ) 
    
    # 查找index
    with open("detail.csv", 'r') as csvfile:
        reader = csv.DictReader(csvfile)
        for row in reader:
            if row['序号'] == str(index):
                pick = row
                #print(row['型号'])

    # 读取本地库
    with open("paragraph.csv", 'r') as csvfile:
        reader = csv.reader(csvfile)
        for row in reader:
            temp_num = 0
            temp_val = []
            for i, val in enumerate(row):
                if val != '':
                    temp_num += 1
                    temp_val.append(val)
                    #print(val)
            r_num.append(temp_num)
            r_val.append(temp_val)
                
#    # 测试是否完全存储         
#    for i, temp_num in enumerate(r_num):
#        for j in r_val[i]:
#            print(j)

#    # 伪随机数生成模块。如果不提供 seed，默认使用系统时间。使用相同的 seed，可以获得完全相同的随机数序列，常用于算法改进测试。
#    random.seed(10086)

    pl1 = []
    pl2 = []
    pl3 = []
    # 生成三段话
    for i, temp_num in enumerate(r_num):
        if i < 3:
            if i == 1:
                if pick['分类'] == 'TFT Digital':
                    pl1.append(r_val[i][0])
                elif pick['分类'] == 'Automotive TFT':
                    pl1.append(r_val[i][1])
                elif pick['分类'] == 'Monochrome-Character':
                    pl1.append(r_val[i][2])
                elif pick['分类'] == 'Monochrome-Graphic':
                    pl1.append(r_val[i][3])
            else:
                r1 = random.randint(0, r_num[i] - 1)
                pl1.append(r_val[i][r1])
        elif i < 7:
            if i == 4:
                if pick['类型'] == 'Transflective':
                    pl2.append(r_val[i][0])
                elif pick['类型'] == 'Transmissive':
                    pl2.append(r_val[i][1])
            else:
                r2 = random.randint(0, r_num[i] - 1)
                pl2.append(r_val[i][r2])
        else:
            r3 = random.randint(0, r_num[i] - 1)
            pl3.append(r_val[i][r3])
            
#    # 测试输出结果        
#    print(pl1)
#    print(pl2)
#    print(pl3)

    temp1 = pl1[0].replace('XXX', pick['型号'])
    
    
    
    temp1 = temp1.replace('YYY', pl1[1])
    p1 = temp1 + pl1[2]

    temp1 = pl2[0].replace('XXX', pick['型号'])
    if pick['大小'] == '':
        err.append('大小')
    else:
        temp1 = temp1.replace('YYY', pick['大小'])
    if pick['分辨率'] == '':
        err.append('分辨率')
    else:
        temp1 = temp1.replace('ZZZ', pick['分辨率'])
    if pick['背光'] == '':
        err.append('背光')
    else:
        temp3 = pl2[2].replace('XXX', pick['背光'])
    p2 = temp1 + pl2[1] + temp3 + pl2[3]

    if pick['介质'] == '':
        err.append('介质')
    else:
        temp1 = pl3[0].replace('XXX', pick['介质'])
    if pick['亮度'] == '':
        err.append('亮度')
    else:
        temp1 = temp1.replace('YYY', pick['亮度'])
    if pick['对比度'] == '':
        err.append('对比度')
    else:
        contrast = pick['对比度']
        contrast_list = contrast.split(':', 1)
        contrast = contrast_list[0] + ':1'
        temp1 = temp1.replace('ZZZ', contrast)
    
    if pick['Operating Temperature (oC)'] == '':
        err.append('操作温度')
    else: 
        temp2 = pl3[1].replace('XXX', pick['Operating Temperature (oC)'])
    if pick['Storage Temperature (oC)'] == '':
        err.append('储存温度')
    else:
        temp2 = temp2.replace('YYY', pick['Storage Temperature (oC)'])
    if pick['Outer Dimension (mm):'] == '':
        err.append('尺寸')
    else:
        temp_d = pick['Outer Dimension (mm):']
        temp_d = temp_d.replace(' ', '')
        temp_d = temp_d.replace('x', ' X ')
        temp3 = pl3[2].replace('XXX', temp_d)
    if pick['Mass (g):'] == '':
        err.append('重量')
    else:
        temp3 = temp3.replace('YYY', pick['Mass (g):'])
    p3 = temp1 + temp2 + temp3 + pl3[3]
    err.append('视角')
    
#    # 测试输出结果 
#    print(p1)
#    print(p2)
#    print(p3)
    
    MyContent_type = '新产品'
    MyContent_factory = 'Kyocera(京瓷)'
    MyContent_key = '液晶显示器，TFT，薄膜晶体管，显示屏'
    MyContent_app = '工业，消费'
    MyContent_category = '改写'
    MyContent_author = '刘阳'

    doc = DocxTemplate("templet.docx")
    context = {
        'MyTable' : [
            {'title' : '【产品】' + pick['标题'], 'type' : MyContent_type, 'factory' : MyContent_factory, 'version' : pick['型号'], 
             'abstract' : '', 'key' : MyContent_key, 'app' : MyContent_app, 'category' : MyContent_category, 
             'ref' : pick['网址'], 'author' : MyContent_author }
        ],
        'MyFirstP' : p1 ,
        'MySecondP' : p2 ,
        'MyThirdP' : p3
    }
    
    
    doc.render(context)
    doc.save("generated_temp.docx")
    
    
    # 源文件 test.docx
    doc = Document("generated_temp.docx")
    #doc = Document("刘阳+PXFC211507SC.docx")
    #styles = doc.styles
    #
    #for style in styles:
    #    print(style.name)
    
    p1 = doc.add_paragraph('')
    r1 = p1.add_run(pick['型号'] + '的主要特点：')
    
    if pick['特性1'] != '':
        paragraph = doc.add_paragraph(pick['特性1'])
        paragraph_format = paragraph.paragraph_format
    if pick['特性2'] != '':
        paragraph = doc.add_paragraph(pick['特性2'])
        paragraph_format = paragraph.paragraph_format
    
    doc.add_paragraph('')
    p2 = doc.add_paragraph('')
    r2 = p2.add_run(pick['型号'] + '的典型应用：')
    
    application = pick['应用']
    application = application.replace('?', '')
    app_list = application.split(',')
    for i in app_list:
        i = i.strip()
        paragraph = doc.add_paragraph(i)  
        paragraph_format = paragraph.paragraph_format
    
    
    file_name = '刘阳+' + pick['型号'] + '.docx'
    
    doc.save(file_name)    

    print(err)
    
    